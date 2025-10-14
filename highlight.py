#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Kindle Ebook Highlighter - COMPLETE VERSION with Batch Processing

All features included:
- Batch processing enabled by DEFAULT (process ALL books with 95%+ match confidence)
- Library mode enabled by DEFAULT
- Text formatting preservation (bold, italic, underline, links, footnotes)
- Proper paragraph structure preservation
- All 3 matching methods: regex, difflib, vector
- Compare mode
- HTML preservation for debugging (--preserve-html)
- Enhanced verbose logging with -vv
- Improved robustness and error handling

Usage:
  Default: Batch process ALL books from library with 95%+ match confidence:
    python kindle_highlighter.py --clippings "My Clippings.txt" -v
  
  Process single best match only:
    python kindle_highlighter.py --clippings "My Clippings.txt" --no-batch -v
  
  Process specific ebook file:
    python kindle_highlighter.py --ebook book.mobi --clippings "My Clippings.txt" -v
  
  Batch with HTML preservation for debugging:
    python kindle_highlighter.py --clippings "My Clippings.txt" --preserve-html -vv
  
  Compare all methods on specific file:
    python kindle_highlighter.py --ebook book.mobi --clippings "My Clippings.txt" --compare -v
  
  List all books in library:
    python kindle_highlighter.py --list-books
"""

import sys
import os
import re
import argparse
import subprocess
import logging
import zipfile
import shutil
import glob
from typing import List, Dict, Tuple, Callable, Optional, Set
from pathlib import Path
from difflib import SequenceMatcher
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt
from bs4 import BeautifulSoup, NavigableString, Tag
from thefuzz import process as fuzzy_process
from thefuzz import fuzz

os.environ['USE_TF'] = '0'
os.environ['USE_TORCH'] = '1'
os.environ['TRANSFORMERS_NO_TF'] = '1'
os.environ['TF_CPP_MIN_LOG_LEVEL'] = '3'  # Suppress TensorFlow warnings

# --- CONFIGURATION ---
DEFAULT_CALIBRE_LIBRARY = os.path.expanduser('~/Calibre Library/')
DEFAULT_SIMILARITY_THRESHOLD = 0.9
DEFAULT_VECTOR_CONFIDENCE_THRESHOLD = 0.65
NLP_MODEL = 'all-MiniLM-L6-v2'
MAX_BOOK_SIZE_MB = 50
MIN_HIGHLIGHT_LENGTH = 5
MIN_TITLE_MATCH_SCORE = 90
AUTO_SELECT_THRESHOLD = 95

# --- LOGGING SETUP ---
logger = logging.getLogger('kindle_highlighter')

def setup_logging(verbosity: int):
    """Configure logging based on verbosity level."""
    if verbosity == 0:
        level = logging.WARNING
    elif verbosity == 1:
        level = logging.INFO
    else:  # 2 or more
        level = logging.DEBUG
    
    handler = logging.StreamHandler(sys.stdout)
    if level == logging.DEBUG:
        formatter = logging.Formatter('%(levelname)s [%(funcName)s]: %(message)s')
    else:
        formatter = logging.Formatter('%(levelname)s: %(message)s')
    handler.setFormatter(formatter)
    logger.addHandler(handler)
    logger.setLevel(level)
    
    logger.debug(f"Logging configured at level: {logging.getLevelName(level)}")

# --- HELPER FUNCTIONS ---

def handle_error(message: str, is_fatal: bool = True) -> None:
    """Prints an error message and optionally exits."""
    logger.error(message)
    if is_fatal:
        sys.exit(1)

def normalize_title_for_matching(title: str) -> str:
    """
    Normalize a title for better matching.
    - Remove author names in parentheses
    - Remove series info in parentheses  
    - Remove extra whitespace
    - Remove special chars
    """
    logger.debug(f"Normalizing title: '{title}'")
    
    # Remove author/series in parentheses at the end
    title_clean = re.sub(r'\s*\([^)]+\)\s*$', '', title)
    
    # Remove special unicode characters
    title_clean = title_clean.replace('¬', '').replace('', '')
    
    # Remove multiple spaces
    title_clean = re.sub(r'\s+', ' ', title_clean).strip()
    
    logger.debug(f"Normalized to: '{title_clean}'")
    return title_clean

def find_calibre_binary() -> Optional[str]:
    """Find Calibre's ebook-convert binary on the system."""
    logger.debug("Searching for Calibre installation...")
    
    # macOS
    if sys.platform == 'darwin':
        logger.debug("Platform: macOS - searching /Applications")
        calibre_apps = glob.glob('/Applications/*alibre*.app')
        logger.debug(f"Found {len(calibre_apps)} potential Calibre apps: {calibre_apps}")
        
        for app_path in calibre_apps:
            binary_path = os.path.join(app_path, 'Contents/MacOS/ebook-convert')
            logger.debug(f"Checking: {binary_path}")
            if os.path.exists(binary_path) and os.access(binary_path, os.X_OK):
                logger.debug(f"✓ Found Calibre binary: {binary_path}")
                return binary_path
        
        common_paths = [
            '/Applications/calibre.app/Contents/MacOS/ebook-convert',
            '/Applications/Calibre.app/Contents/MacOS/ebook-convert',
            '/usr/local/bin/ebook-convert',
        ]
        logger.debug(f"Checking common macOS paths: {common_paths}")
        for path in common_paths:
            if os.path.exists(path) and os.access(path, os.X_OK):
                logger.debug(f"✓ Found Calibre binary: {path}")
                return path
    
    # Linux/Unix
    elif sys.platform.startswith('linux') or sys.platform.startswith('freebsd'):
        logger.debug(f"Platform: {sys.platform} - checking Linux/Unix paths")
        common_paths = [
            '/usr/bin/ebook-convert',
            '/usr/local/bin/ebook-convert',
            os.path.expanduser('~/calibre/ebook-convert'),
        ]
        logger.debug(f"Checking paths: {common_paths}")
        for path in common_paths:
            if os.path.exists(path) and os.access(path, os.X_OK):
                logger.debug(f"✓ Found Calibre binary: {path}")
                return path
    
    # Windows
    elif sys.platform == 'win32':
        logger.debug("Platform: Windows - checking Program Files")
        common_paths = [
            os.path.join(os.environ.get("ProgramFiles", "C:\\Program Files"), "Calibre2", "ebook-convert.exe"),
            os.path.join(os.environ.get("ProgramFiles(x86)", "C:\\Program Files (x86)"), "Calibre2", "ebook-convert.exe"),
            os.path.join(os.environ.get("ProgramFiles", "C:\\Program Files"), "Calibre", "ebook-convert.exe"),
            os.path.join(os.environ.get("ProgramFiles(x86)", "C:\\Program Files (x86)"), "Calibre", "ebook-convert.exe"),
        ]
        logger.debug(f"Checking paths: {common_paths}")
        for path in common_paths:
            if os.path.exists(path):
                logger.debug(f"✓ Found Calibre binary: {path}")
                return path
    
    # Try in PATH
    logger.debug("Checking system PATH")
    try:
        command = 'where' if sys.platform == 'win32' else 'which'
        logger.debug(f"Running: {command} ebook-convert")
        result = subprocess.run(
            [command, 'ebook-convert'],
            capture_output=True,
            text=True,
            timeout=5
        )
        if result.returncode == 0:
            path = result.stdout.strip().split('\n')[0]
            logger.debug(f"✓ Found ebook-convert in PATH: {path}")
            return path
        else:
            logger.debug(f"Command failed with return code: {result.returncode}")
    except (FileNotFoundError, subprocess.TimeoutExpired) as e:
        logger.debug(f"PATH search failed: {e}")
    except Exception as e:
        logger.debug(f"Unexpected error in PATH search: {e}")
    
    logger.debug("Calibre binary not found")
    return None

def check_converter_available(calibre_path: Optional[str] = None) -> Tuple[Optional[str], Optional[str]]:
    """Check which ebook converter is available."""
    logger.debug("Checking for available ebook converters...")
    
    if calibre_path:
        logger.debug(f"Checking custom Calibre path: {calibre_path}")
        if os.path.exists(calibre_path) and os.access(calibre_path, os.X_OK):
            try:
                result = subprocess.run(
                    [calibre_path, '--version'],
                    capture_output=True,
                    text=True,
                    timeout=5
                )
                if result.returncode == 0:
                    version = result.stdout.strip()
                    logger.debug(f"Custom Calibre version: {version}")
                    logger.info(f"Using Calibre (custom path): {calibre_path}")
                    return calibre_path, 'htmlz'
            except (subprocess.TimeoutExpired, Exception) as e:
                logger.debug(f"Custom Calibre path failed: {e}")
    
    logger.debug("Auto-detecting Calibre...")
    calibre_binary = find_calibre_binary()
    if calibre_binary:
        try:
            result = subprocess.run(
                [calibre_binary, '--version'],
                capture_output=True,
                text=True,
                timeout=5
            )
            if result.returncode == 0:
                version = result.stdout.strip()
                logger.debug(f"Calibre version: {version}")
                logger.info(f"Using Calibre: {calibre_binary}")
                return calibre_binary, 'htmlz'
        except (subprocess.TimeoutExpired, Exception) as e:
            logger.debug(f"Calibre detection failed: {e}")
    
    logger.debug("Checking for standalone ebook-converter...")
    try:
        result = subprocess.run(
            ['ebook-converter', '--help'],
            capture_output=True,
            text=True,
            timeout=5
        )
        if result.returncode == 0:
            logger.debug("✓ Found standalone ebook-converter")
            logger.info("Using standalone ebook-converter")
            return 'ebook-converter', 'htmlz'
    except (FileNotFoundError, subprocess.TimeoutExpired) as e:
        logger.debug(f"Standalone ebook-converter not found: {e}")
    except Exception as e:
        logger.debug(f"Unexpected error checking ebook-converter: {e}")
    
    logger.warning("No ebook converter found")
    return None, None

def check_dependencies(calibre_path: Optional[str] = None) -> Dict[str, bool]:
    """Check all required and optional dependencies."""
    logger.debug("Checking Python package dependencies...")
    deps = {
        'docx': False,
        'bs4': False,
        'thefuzz': False,
        'converter': False,
        'mobi': False,
        'ebooklib': False
    }
    
    try:
        import docx
        deps['docx'] = True
        logger.debug("✓ python-docx found")
    except ImportError:
        logger.error("✗ python-docx not found. Install: pip install python-docx")
    
    try:
        import bs4
        deps['bs4'] = True
        logger.debug("✓ beautifulsoup4 found")
    except ImportError:
        logger.error("✗ beautifulsoup4 not found. Install: pip install beautifulsoup4")
    
    try:
        import thefuzz
        deps['thefuzz'] = True
        logger.debug("✓ thefuzz found")
    except ImportError:
        logger.error("✗ thefuzz not found. Install: pip install thefuzz python-Levenshtein")
    
    logger.debug("Checking for ebook converter...")
    converter, _ = check_converter_available(calibre_path)
    deps['converter'] = converter is not None
    if not deps['converter']:
        logger.error("✗ No ebook converter found. Install Calibre or ebook-converter")
    else:
        logger.debug("✓ Ebook converter found")
    
    logger.debug("Checking optional conversion libraries...")
    try:
        import mobi
        deps['mobi'] = True
        logger.debug("✓ mobi library found (optional)")
    except ImportError:
        logger.debug("mobi library not found (optional: pip install mobi)")
    
    try:
        import ebooklib
        deps['ebooklib'] = True
        logger.debug("✓ ebooklib found (optional)")
    except ImportError:
        logger.debug("ebooklib not found (optional: pip install ebooklib)")
    
    return deps

def check_nlp_libraries() -> bool:
    """Checks if the required NLP libraries are installed."""
    logger.debug("Checking NLP libraries for vector matching...")
    
    # Disable TensorFlow to avoid import conflicts
    import os
    os.environ['USE_TF'] = '0'
    os.environ['USE_TORCH'] = '1'
    os.environ['TRANSFORMERS_NO_TF'] = '1'
    
    try:
        import torch
        logger.debug("✓ torch found")
    except ImportError as e:
        logger.debug(f"torch not found: {e}")
        return False
    
    try:
        import nltk
        logger.debug("✓ nltk found")
    except ImportError as e:
        logger.debug(f"nltk not found: {e}")
        return False
    
    try:
        # Try importing with TF disabled
        from sentence_transformers import SentenceTransformer, util
        logger.debug("✓ sentence_transformers found")
        return True
    except ImportError as e:
        logger.debug(f"sentence_transformers not found: {e}")
        return False
    except RuntimeError as e:
        error_msg = str(e)
        if 'tensorflow' in error_msg.lower() or 'tf_keras' in error_msg.lower():
            logger.error("=" * 70)
            logger.error("TensorFlow compatibility issue detected!")
            logger.error("=" * 70)
            logger.error("Your environment has a TensorFlow/transformers version conflict.")
            logger.error("\nTo fix this, run ONE of these solutions:")
            logger.error("\n1. Uninstall TensorFlow (sentence_transformers only needs PyTorch):")
            logger.error("   pip uninstall tensorflow tf-keras")
            logger.error("\n2. OR create a fresh environment:")
            logger.error("   conda create -n highlighter python=3.11")
            logger.error("   conda activate highlighter")
            logger.error("   pip install sentence-transformers torch nltk")
            logger.error("\n3. OR use hybrid/regex/diff methods instead (no NLP needed):")
            logger.error("   python highlight.py --clippings ... -m hybrid")
            logger.error("=" * 70)
            return False
        else:
            logger.debug(f"Unexpected runtime error: {e}")
            return False
    except Exception as e:
        logger.debug(f"Unexpected error checking NLP libraries: {e}")
        return False

# --- CALIBRE LIBRARY FUNCTIONS ---

def find_calibre_library(library_path: Optional[str] = None) -> Path:
    """Find and validate Calibre library path."""
    if library_path:
        path = Path(library_path).expanduser()
        logger.debug(f"Using provided library path: {path}")
    else:
        path = Path(DEFAULT_CALIBRE_LIBRARY).expanduser()
        logger.debug(f"Using default library path: {path}")
    
    if not path.exists():
        handle_error(f"Calibre library not found at: {path}")
    
    if not path.is_dir():
        handle_error(f"Calibre library path is not a directory: {path}")
    
    logger.info(f"Using Calibre library: {path}")
    return path

def discover_books(library_path: Path, format_filter: str = 'mobi') -> List[Dict]:
    """Recursively discover all books in Calibre library."""
    logger.info(f"Scanning library for .{format_filter} files...")
    logger.debug(f"Starting recursive scan of: {library_path}")
    
    books = []
    format_filter = format_filter.lower()
    scanned_dirs = 0
    scanned_files = 0
    
    try:
        for root, dirs, files in os.walk(library_path):
            scanned_dirs += 1
            if scanned_dirs % 100 == 0:
                logger.debug(f"Scanned {scanned_dirs} directories, found {len(books)} books so far...")
            
            for file in files:
                scanned_files += 1
                if file.lower().endswith(f'.{format_filter}'):
                    file_path = Path(root) / file
                    logger.debug(f"Found ebook: {file_path}")
                    
                    metadata_file = Path(root) / 'metadata.opf'
                    title = None
                    author = None
                    
                    if metadata_file.exists():
                        logger.debug(f"Reading metadata from: {metadata_file}")
                        try:
                            with open(metadata_file, 'r', encoding='utf-8') as f:
                                soup = BeautifulSoup(f, 'xml')
                                title_tag = soup.find('dc:title')
                                author_tag = soup.find('dc:creator')
                                if title_tag:
                                    title = title_tag.get_text(strip=True)
                                    logger.debug(f"  Title from metadata: {title}")
                                if author_tag:
                                    author = author_tag.get_text(strip=True)
                                    logger.debug(f"  Author from metadata: {author}")
                        except Exception as e:
                            logger.debug(f"Failed to parse metadata: {e}")
                    
                    if not title:
                        folder_name = Path(root).name
                        title = re.sub(r'\s*\(\d+\)$', '', folder_name)
                        logger.debug(f"  Using folder name as title: {title}")
                    
                    books.append({
                        'path': file_path,
                        'title': title or file,
                        'author': author or 'Unknown',
                        'filename': file,
                        'folder': Path(root).name
                    })
    except Exception as e:
        logger.error(f"Error during library scan: {e}")
        logger.debug("Full traceback:", exc_info=True)
    
    logger.info(f"✓ Found {len(books)} books in library (scanned {scanned_dirs} directories, {scanned_files} files)")
    if books:
        logger.debug(f"First 10 books: {[b['title'] for b in books[:10]]}")
    
    return books

def match_books_to_clippings(books: List[Dict], all_highlights: List[Dict[str, str]], all_notes: List[Dict[str, str]], min_score: int = MIN_TITLE_MATCH_SCORE) -> List[Dict]:
    """
    Match ALL books from library to clippings.
    Returns list of matched books sorted by score and highlight count.
    """
    if not books:
        logger.warning("No books found in library")
        return []
    
    # Combine highlights and notes to get all clipping titles
    clippings_titles = [h['title'] for h in all_highlights] + [n['title'] for n in all_notes]
    
    if not clippings_titles:
        logger.warning("No clipping titles to match against")
        return []
    
    from collections import Counter
    title_counts = Counter(clippings_titles)
    
    logger.info(f"Matching {len(books)} library books against {len(title_counts)} unique titles from clippings...")
    logger.debug(f"Clipping titles: {list(title_counts.keys())[:10]}{'...' if len(title_counts) > 10 else ''}")
    
    matches = []
    processed_titles = 0
    
    for clip_title, highlight_count in title_counts.most_common():
        processed_titles += 1
        if processed_titles % 10 == 0:
            logger.debug(f"Processing clipping title {processed_titles}/{len(title_counts)}")
        
        clip_title_clean = normalize_title_for_matching(clip_title)
        
        best_score = 0
        best_book = None
        best_method = None
        
        for book in books:
            book_title_clean = normalize_title_for_matching(book['title'])
            
            score1 = fuzz.token_sort_ratio(clip_title_clean.lower(), book_title_clean.lower())
            score2 = fuzz.partial_ratio(clip_title_clean.lower(), book_title_clean.lower())
            score3 = fuzz.ratio(clip_title_clean.lower(), book_title_clean.lower())
            
            score = max(score1, score2, score3)
            
            if score > best_score:
                best_score = score
                best_book = book
                best_method = f"s1={score1},s2={score2},s3={score3}"
        
        if best_book and best_score >= min_score:
            matches.append({
                'book': best_book,
                'clipping_title': clip_title,
                'clipping_title_normalized': clip_title_clean,
                'highlight_count': highlight_count,
                'score': best_score,
                'method_scores': best_method
            })
            
            logger.debug(f"'{clip_title_clean}' ({highlight_count} items) -> "
                        f"'{normalize_title_for_matching(best_book['title'])}' ({best_score}%) [{best_method}]")
    
    matches.sort(key=lambda x: (x['score'], x['highlight_count']), reverse=True)
    
    logger.debug(f"Match complete: {len(matches)} books matched out of {len(title_counts)} clipping titles")
    return matches

# --- FILE PARSING ---

def validate_file_size(filepath: str, max_size_mb: int = MAX_BOOK_SIZE_MB) -> bool:
    """Validate that file size is reasonable."""
    try:
        size_mb = os.path.getsize(filepath) / (1024 * 1024)
        logger.debug(f"File size: {size_mb:.2f}MB")
        if size_mb > max_size_mb:
            logger.warning(f"File is large ({size_mb:.2f}MB). This may take a while...")
        return True
    except Exception as e:
        logger.error(f"Could not check file size: {e}")
        return False

def parse_clippings(clippings_file: str) -> Tuple[List[Dict[str, str]], List[Dict[str, str]]]:
    """
    Parses the 'My Clippings.txt' file.
    Returns: (highlights, notes) as separate lists
    """
    if not os.path.exists(clippings_file):
        handle_error(f"Clippings file not found at: {clippings_file}")
    
    logger.info(f"Parsing clippings from: {clippings_file}")
    
    try:
        file_size = os.path.getsize(clippings_file) / (1024 * 1024)
        logger.debug(f"Clippings file size: {file_size:.2f}MB")
        
        if file_size > 10:
            logger.info(f"Large clippings file ({file_size:.2f}MB). This may take a moment...")
    except Exception as e:
        logger.debug(f"Could not check file size: {e}")
    
    encodings = ['utf-8-sig', 'utf-8', 'latin-1', 'cp1252']
    content = None
    
    for encoding in encodings:
        logger.debug(f"Trying encoding: {encoding}")
        try:
            with open(clippings_file, 'r', encoding=encoding) as f:
                content = f.read()
            logger.debug(f"✓ Successfully read file with {encoding} encoding")
            break
        except UnicodeDecodeError:
            logger.debug(f"✗ Failed to read with {encoding} encoding")
            continue
        except Exception as e:
            logger.debug(f"Unexpected error with {encoding}: {e}")
            continue
    
    if content is None:
        handle_error(f"Could not read clippings file with any known encoding")

    logger.debug(f"File content length: {len(content)} characters")
    logger.info("Parsing clippings (this may take a moment for large files)...")
    
    highlights = []
    notes = []
    entries = content.split('==========')
    logger.debug(f"Found {len(entries)} potential entries (split by '==========')")
    
    highlight_patterns = [
        r'^-\s*Your\s+Highlight',
        r'^-\s*(Ihre|Deine)\s+Markierung',
        r'^-\s*Votre\s+surlignement',
        r'^-\s*Tu\s+subrayado',
        r'^-\s*La\s+tua\s+evidenziazione',
    ]
    
    note_patterns = [
        r'^-\s*Your\s+Note',
        r'^-\s*(Ihre|Deine)\s+Notiz',
        r'^-\s*Votre\s+note',
        r'^-\s*Tu\s+nota',
        r'^-\s*La\s+tua\s+nota',
    ]
    
    logger.debug(f"Using {len(highlight_patterns)} highlight patterns and {len(note_patterns)} note patterns")
    
    # Extract location info pattern (works for most Kindle formats)
    # Example: "- Your Highlight on page 52 | location 789-791 | Added on Wednesday, 1 January 2025 12:00:00"
    location_pattern = r'(?:page\s+(\d+)|location\s+(\d+)(?:-(\d+))?|Loc\.\s+(\d+)(?:-(\d+))?)'
    
    skipped_no_lines = 0
    skipped_not_recognized = 0
    skipped_no_text = 0
    skipped_too_short = 0
    
    for idx, entry in enumerate(entries):
        if idx % 100 == 0 and idx > 0:
            logger.debug(f"Processing entry {idx}/{len(entries)}")
        
        entry = entry.strip()
        if not entry:
            continue
        
        lines = entry.split('\n')
        if len(lines) < 3:
            skipped_no_lines += 1
            if idx < 10:
                logger.debug(f"Skipping entry {idx}: too few lines ({len(lines)})")
            continue
        
        book_title = ''.join(char for char in lines[0].strip() if char.isprintable())
        
        # Extract location information
        location_info = {}
        location_match = re.search(location_pattern, lines[1], re.IGNORECASE)
        if location_match:
            groups = location_match.groups()
            if groups[0]:  # page number
                location_info['page'] = int(groups[0])
            if groups[1]:  # location start
                location_info['location_start'] = int(groups[1])
            if groups[2]:  # location end
                location_info['location_end'] = int(groups[2])
            if groups[3]:  # alternative Loc. format start
                location_info['location_start'] = int(groups[3])
            if groups[4]:  # alternative Loc. format end
                location_info['location_end'] = int(groups[4])
        
        # Check if it's a note
        is_note = False
        for pattern in note_patterns:
            if re.match(pattern, lines[1].strip(), re.IGNORECASE):
                is_note = True
                break
        
        # Check if it's a highlight
        is_highlight = False
        for pattern in highlight_patterns:
            if re.match(pattern, lines[1].strip(), re.IGNORECASE):
                is_highlight = True
                break
        
        if not is_highlight and not is_note:
            skipped_not_recognized += 1
            if idx < 10:
                logger.debug(f"Skipping entry {idx}: not recognized as highlight or note: {lines[1][:80]}")
            continue
        
        # Extract text
        clipping_text = None
        for i in range(2, len(lines)):
            if lines[i].strip():
                clipping_text = '\n'.join(lines[i:]).strip()
                break
        
        if not clipping_text:
            skipped_no_text += 1
            logger.debug(f"Skipping entry {idx}: no text found")
            continue
        
        # Add to appropriate list
        if len(clipping_text) >= MIN_HIGHLIGHT_LENGTH or is_note:  # Notes can be shorter
            item = {
                'title': book_title,
                'text': clipping_text,
                'location': location_info
            }
            
            if is_highlight:
                highlights.append(item)
                if idx < 10:
                    logger.debug(f"✓ Added highlight from '{book_title}': {clipping_text[:50]}...")
            else:  # is_note
                notes.append(item)
                if idx < 10:
                    logger.debug(f"✓ Added note from '{book_title}': {clipping_text[:50]}...")
        else:
            skipped_too_short += 1
            logger.debug(f"Skipping entry {idx}: text too short ({len(clipping_text)} chars)")
    
    logger.debug(f"Parsing stats: skipped {skipped_no_lines} (no lines), "
                 f"{skipped_not_recognized} (not recognized), {skipped_no_text} (no text), {skipped_too_short} (too short)")
    
    if not highlights and not notes:
        logger.error("No valid clippings were found in the provided file.")
        logger.error("The file may not be in the standard Kindle 'My Clippings.txt' format.")
        handle_error("Failed to parse clippings file.")
    
    unique_highlight_titles = len(set(h['title'] for h in highlights))
    unique_note_titles = len(set(n['title'] for n in notes))
    
    logger.info(f"✓ Found {len(highlights)} highlights from {unique_highlight_titles} books")
    logger.info(f"✓ Found {len(notes)} user notes from {unique_note_titles} books")
    
    return highlights, notes

# --- CONVERSION FUNCTIONS ---

def extract_htmlz(htmlz_path: str) -> Tuple[str, str]:
    """Extract htmlz file and return paths to main HTML and extract dir."""
    logger.debug(f"Extracting htmlz file: {htmlz_path}")
    
    extract_dir = os.path.splitext(htmlz_path)[0] + '_extracted'
    logger.debug(f"Extract directory: {extract_dir}")
    
    if os.path.exists(extract_dir):
        logger.debug(f"Removing existing extract directory")
        try:
            shutil.rmtree(extract_dir)
        except Exception as e:
            logger.warning(f"Could not remove existing extract directory: {e}")
    
    try:
        logger.debug("Opening zip archive...")
        with zipfile.ZipFile(htmlz_path, 'r') as zip_ref:
            file_list = zip_ref.namelist()
            logger.debug(f"Archive contains {len(file_list)} files")
            if logger.level <= logging.DEBUG and len(file_list) <= 20:
                logger.debug(f"Files in archive: {file_list}")
            zip_ref.extractall(extract_dir)
        
        html_files = list(Path(extract_dir).rglob('*.html'))
        logger.debug(f"Found {len(html_files)} HTML files in archive")
        
        if not html_files:
            handle_error(f"No HTML files found in {htmlz_path}")
        
        main_html = None
        for html_file in html_files:
            logger.debug(f"Checking HTML file: {html_file.name}")
            if html_file.name.lower() in ['index.html', 'index.htm']:
                main_html = html_file
                logger.debug(f"✓ Found index file: {main_html}")
                break
        
        if not main_html:
            main_html = html_files[0]
            logger.debug(f"Using first HTML file: {main_html}")
        
        logger.debug(f"✓ Selected HTML file: {main_html}")
        return str(main_html), extract_dir
    
    except zipfile.BadZipFile:
        handle_error(f"Failed to extract {htmlz_path}. Not a valid zip file.")
    except Exception as e:
        logger.error(f"Error during extraction: {e}")
        logger.debug("Full traceback:", exc_info=True)
        handle_error(f"Failed to extract htmlz: {e}")

def convert_ebook_to_html(ebook_path: str, converter_path: str, output_format: str) -> Tuple[str, Optional[str], Optional[str]]:
    """
    Converts an ebook to HTML.
    Returns: (html_file_path, intermediate_file, extract_dir)
    """
    if not os.path.exists(ebook_path):
        handle_error(f"Ebook file not found at: {ebook_path}")
    
    logger.debug(f"Starting conversion of: {ebook_path}")
    validate_file_size(ebook_path)
    
    if output_format == 'htmlz':
        output_path = os.path.splitext(ebook_path)[0] + '.htmlz'
    else:
        output_path = os.path.splitext(ebook_path)[0] + '.html'
    
    logger.debug(f"Output format: {output_format}")
    logger.debug(f"Output path: {output_path}")
    logger.info(f"Converting '{os.path.basename(ebook_path)}' to {output_format.upper()}...")
    
    try:
        logger.debug(f"Running converter: {converter_path}")
        logger.debug(f"Command: {converter_path} {ebook_path} {output_path}")
        
        result = subprocess.run(
            [converter_path, ebook_path, output_path],
            check=True,
            capture_output=True,
            text=True,
            timeout=300
        )
        
        if logger.level <= logging.DEBUG:
            if result.stdout:
                logger.debug(f"Converter stdout: {result.stdout[:500]}...")
            if result.stderr:
                logger.debug(f"Converter stderr: {result.stderr[:500]}...")
        
        logger.info("✓ Conversion successful")
        
        if not os.path.exists(output_path):
            handle_error("Output file was not created by converter")
        
        output_size = os.path.getsize(output_path)
        logger.debug(f"Output file size: {output_size} bytes ({output_size/1024:.1f} KB)")
        
        if output_size < 100:
            handle_error("Generated file is suspiciously small. Conversion may have failed.")
        
        if output_format == 'htmlz':
            logger.debug("Extracting htmlz archive...")
            html_file, extract_dir = extract_htmlz(output_path)
            return html_file, output_path, extract_dir
        
        return output_path, None, None
    
    except FileNotFoundError:
        handle_error(f"The command '{converter_path}' was not found.")
    except subprocess.TimeoutExpired:
        handle_error("Conversion timed out. The file may be too large or corrupted.")
    except subprocess.CalledProcessError as e:
        logger.error(f"Converter failed with error:\n{e.stderr}")
        logger.debug("Full error details:", exc_info=True)
        handle_error("Conversion failed. Check that the file is not DRM-protected.")
    except Exception as e:
        logger.error(f"Unexpected error during conversion: {e}")
        logger.debug("Full traceback:", exc_info=True)
        handle_error(f"Conversion failed: {e}")

def try_native_conversion(ebook_path: str) -> Tuple[Optional[str], None, None]:
    """Attempt to convert using native Python libraries as fallback."""
    logger.info("Attempting native Python conversion...")
    logger.debug(f"Trying to convert: {ebook_path}")
    
    try:
        logger.debug("Trying mobi library...")
        import mobi
        
        tempdir, filepath = mobi.extract(ebook_path)
        logger.debug(f"mobi extracted to: {tempdir}")
        logger.debug(f"mobi HTML file: {filepath}")
        logger.info(f"✓ Extracted with mobi library: {filepath}")
        return filepath, None, None
    
    except ImportError:
        logger.debug("mobi library not available (install: pip install mobi)")
    except Exception as e:
        logger.debug(f"mobi library failed: {e}")
        logger.debug("mobi exception details:", exc_info=True)
    
    try:
        logger.debug("Trying ebooklib...")
        import ebooklib
        from ebooklib import epub
        
        book = epub.read_epub(ebook_path)
        html_content = []
        
        items = list(book.get_items())
        logger.debug(f"Found {len(items)} items in epub")
        
        for item in items:
            if item.get_type() == ebooklib.ITEM_DOCUMENT:
                html_content.append(item.get_content().decode('utf-8'))
        
        logger.debug(f"Extracted {len(html_content)} HTML documents")
        
        output_path = os.path.splitext(ebook_path)[0] + '_native.html'
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('<html><body>')
            f.write(''.join(html_content))
            f.write('</body></html>')
        
        logger.debug(f"Wrote combined HTML to: {output_path}")
        logger.info(f"✓ Converted with ebooklib: {output_path}")
        return output_path, None, None
    
    except ImportError:
        logger.debug("ebooklib not available (install: pip install ebooklib)")
    except Exception as e:
        logger.debug(f"ebooklib conversion failed: {e}")
        logger.debug("ebooklib exception details:", exc_info=True)
    
    logger.error("All native conversion methods failed")
    logger.error("Install native converters: pip install mobi ebooklib")
    return None, None, None

def find_book_highlights_and_notes(ebook_title: str, all_highlights: List[Dict[str, str]], all_notes: List[Dict[str, str]]) -> Tuple[List[str], List[Dict[str, str]]]:
    """Uses fuzzy matching to find highlights and notes for the specified ebook."""
    highlight_titles = list(set(h['title'] for h in all_highlights))
    note_titles = list(set(n['title'] for n in all_notes))
    
    all_titles = list(set(highlight_titles + note_titles))
    
    logger.debug(f"Searching among {len(all_titles)} unique book titles")
    logger.debug(f"Ebook title to match: '{ebook_title}'")
    
    result = fuzzy_process.extractOne(ebook_title, all_titles)
    if result is None:
        logger.warning("Could not match ebook to any clippings titles.")
        return [], []
    
    best_match, score = result
    
    logger.info(f"📖 Ebook Title: '{ebook_title}'")
    logger.info(f"🎯 Best Match from Clippings: '{best_match}' (Confidence: {score}%)")
    
    if score < 75:
        logger.warning("Low confidence in title match.")
        logger.warning("The highlights/notes may be from a different edition or book.")

    # Get highlights
    highlights = [h['text'] for h in all_highlights if h['title'] == best_match]
    
    # Get notes with location info preserved
    book_notes = [n for n in all_notes if n['title'] == best_match]
    
    # Remove duplicates from highlights
    seen = set()
    unique_highlights = []
    for h in highlights:
        if h not in seen:
            seen.add(h)
            unique_highlights.append(h)
    
    if len(highlights) != len(unique_highlights):
        logger.debug(f"Removed {len(highlights) - len(unique_highlights)} duplicate highlights")
    
    logger.debug(f"Returning {len(unique_highlights)} unique highlights and {len(book_notes)} notes")
    return unique_highlights, book_notes

# --- TEXT EXTRACTION WITH PROPER PARAGRAPH STRUCTURE ---

class Paragraph:
    """Represents a paragraph with text and formatting."""
    def __init__(self):
        self.text = ""
        self.formatting = []  # List of (start, end, format_type, format_data)
    
    def add_text(self, text: str, bold: bool = False, italic: bool = False, underline: bool = False, link: str = None):
        """Add text with optional formatting."""
        start = len(self.text)
        self.text += text
        end = len(self.text)
        
        if bold:
            self.formatting.append((start, end, 'bold', None))
        if italic:
            self.formatting.append((start, end, 'italic', None))
        if underline:
            self.formatting.append((start, end, 'underline', None))
        if link:
            self.formatting.append((start, end, 'link', link))

def extract_paragraphs_from_html(soup: BeautifulSoup) -> List[Paragraph]:
    """Extract paragraphs with preserved formatting from HTML."""
    logger.debug("Extracting paragraphs with formatting from HTML...")
    
    paragraphs = []
    
    # Find main content
    main_content = soup.find('body')
    if not main_content:
        main_content = soup
        logger.debug("No <body> tag found, using whole document")
    
    # Find all block-level elements that should be paragraphs
    block_elements = main_content.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li', 'blockquote'])
    
    logger.debug(f"Found {len(block_elements)} potential paragraph elements")
    
    processed_count = 0
    for block_elem in block_elements:
        processed_count += 1
        if processed_count % 100 == 0:
            logger.debug(f"Processing paragraph {processed_count}/{len(block_elements)}")
        
        para = Paragraph()
        
        # Process all children recursively
        def process_element(elem, bold=False, italic=False, underline=False, link=None):
            """Recursively process element and its children."""
            if isinstance(elem, NavigableString):
                text = str(elem)
                # FIXED: Add text even if it's just whitespace (important for spacing!)
                # Only skip if it's completely empty (length 0)
                if len(text) > 0:
                    para.add_text(text, bold=bold, italic=italic, underline=underline, link=link)
            elif isinstance(elem, Tag):
                # Check if this tag adds formatting
                is_bold = bold or elem.name in ['b', 'strong']
                is_italic = italic or elem.name in ['i', 'em']
                is_underline = underline or elem.name == 'u'
                current_link = link or (elem.get('href') if elem.name == 'a' else None)
                
                # Process children
                for child in elem.children:
                    process_element(child, is_bold, is_italic, is_underline, current_link)
        
        # Process the block element
        for child in block_elem.children:
            process_element(child)
        
        # Only add paragraph if it has non-whitespace text
        if para.text.strip():
            paragraphs.append(para)
            if processed_count <= 5:
                logger.debug(f"Paragraph {processed_count}: '{para.text[:80]}...' ({len(para.formatting)} formatting spans)")
    
    logger.debug(f"✓ Extracted {len(paragraphs)} paragraphs with formatting")
    
    if logger.level <= logging.DEBUG and paragraphs:
        total_formatting = sum(len(p.formatting) for p in paragraphs)
        logger.debug(f"Total formatting spans across all paragraphs: {total_formatting}")
    
    return paragraphs

def create_highlighted_docx(
    html_path: str,
    highlights: List[str],
    notes: List[Dict[str, str]],
    doc_title: str,
    matcher_func: Callable,
    threshold: float
) -> Tuple[Document, int]:
    """Create a docx file with highlights and user notes."""
    logger.info(f"🔍 Searching for {len(highlights)} highlights using '{matcher_func.__name__}' method...")
    if notes:
        logger.info(f"📝 Will add {len(notes)} user notes as comments")
    logger.debug(f"HTML path: {html_path}")
    logger.debug(f"Threshold: {threshold}")
    
    logger.debug("Reading and parsing HTML file...")
    try:
        with open(html_path, 'r', encoding='utf-8') as f:
            soup = BeautifulSoup(f, 'html.parser')
    except Exception as e:
        logger.error(f"Failed to read HTML file: {e}")
        logger.debug("Full traceback:", exc_info=True)
        handle_error(f"Could not read HTML file: {e}")

    logger.debug("Extracting paragraphs with formatting...")
    paragraphs = extract_paragraphs_from_html(soup)
    
    if not paragraphs:
        handle_error("No paragraphs extracted from HTML. The book may be DRM-protected or corrupted.")
    
    # Build full text for matching
    logger.debug("Building full text for highlight matching...")
    full_text_parts = []
    para_boundaries = [0]
    
    for para in paragraphs:
        full_text_parts.append(para.text)
        para_boundaries.append(len(''.join(full_text_parts)))
    
    full_text = ''.join(full_text_parts)
    
    logger.debug(f"Full text length: {len(full_text)} characters across {len(paragraphs)} paragraphs")
    logger.debug(f"First 200 chars: {full_text[:200]}...")
    
    logger.debug("Starting highlight matching...")
    found_spans, highlights_found = matcher_func(full_text, highlights, threshold)
    
    percentage = (highlights_found/len(highlights)*100) if len(highlights) > 0 else 0
    logger.info(f"📊 Found {highlights_found} of {len(highlights)} highlights ({percentage:.1f}%)")
    
    if highlights_found == 0 and len(highlights) > 0:
        logger.warning("⚠️  No highlights were found! This could mean:")
        logger.warning("  - The book text was reformatted during conversion")
        logger.warning("  - The highlights are from a different edition")
        logger.warning("  - Try using the 'vector' method: -m vector")
    elif percentage < 50 and len(highlights) > 0:
        logger.warning(f"⚠️  Low match rate ({percentage:.1f}%). Consider using -m vector for better results.")
    
    logger.debug(f"Sorting {len(found_spans)} found spans...")
    found_spans.sort()
    
    # Create set of highlighted character positions
    highlighted_chars = set()
    for start, end in found_spans:
        highlighted_chars.update(range(start, end))
    
    # Try to locate notes in the text (for adding as comments)
    note_positions = []
    if notes:
        logger.debug(f"Attempting to locate {len(notes)} notes in text...")
        for note_idx, note in enumerate(notes):
            # Try to find note context in text using location info or surrounding text
            # For now, we'll collect them and add at the end or as a separate section
            # In the future, we could use location info to place them more precisely
            note_positions.append({
                'text': note['text'],
                'location': note.get('location', {}),
                'position': None  # Could try to estimate position from location info
            })
    
    logger.debug("Creating Word document with proper paragraph structure...")
    doc = Document()
    doc.add_heading(doc_title, level=1)
    
    # Map character positions to paragraph indices
    def char_to_para(pos):
        """Find which paragraph a character position belongs to."""
        for i in range(len(para_boundaries) - 1):
            if para_boundaries[i] <= pos < para_boundaries[i + 1]:
                return i, pos - para_boundaries[i]
        return len(paragraphs) - 1, pos - para_boundaries[-2]
    
    logger.debug("Adding paragraphs to document with highlights and formatting...")
    
    for para_idx, para in enumerate(paragraphs):
        if para_idx % 100 == 0 and para_idx > 0:
            logger.debug(f"Processing paragraph {para_idx}/{len(paragraphs)}")
        
        docx_para = doc.add_paragraph()
        
        # Get global position of this paragraph
        para_start_global = para_boundaries[para_idx]
        
        # Build formatting map for this paragraph
        format_map = {}
        for start, end, fmt_type, fmt_data in para.formatting:
            for pos in range(start, end):
                if pos not in format_map:
                    format_map[pos] = []
                format_map[pos].append((fmt_type, fmt_data))
        
        # Process character by character
        current_run_text = []
        current_run_highlighted = False
        current_run_formats = set()
        
        for local_pos, char in enumerate(para.text):
            global_pos = para_start_global + local_pos
            is_highlighted = global_pos in highlighted_chars
            
            # Get formatting for this character
            char_formats = set()
            if local_pos in format_map:
                for fmt_type, fmt_data in format_map[local_pos]:
                    char_formats.add((fmt_type, fmt_data))
            
            # If formatting or highlight status changes, flush current run
            if current_run_text and (is_highlighted != current_run_highlighted or char_formats != current_run_formats):
                text = ''.join(current_run_text)
                run = docx_para.add_run(text)
                
                if current_run_highlighted:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                
                for fmt_type, fmt_data in current_run_formats:
                    if fmt_type == 'bold':
                        run.bold = True
                    elif fmt_type == 'italic':
                        run.italic = True
                    elif fmt_type == 'underline':
                        run.underline = True
                
                current_run_text = []
            
            current_run_text.append(char)
            current_run_highlighted = is_highlighted
            current_run_formats = char_formats
        
        # Flush final run
        if current_run_text:
            text = ''.join(current_run_text)
            run = docx_para.add_run(text)
            
            if current_run_highlighted:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            
            for fmt_type, fmt_data in current_run_formats:
                if fmt_type == 'bold':
                    run.bold = True
                elif fmt_type == 'italic':
                    run.italic = True
                elif fmt_type == 'underline':
                    run.underline = True
    
    # Add user notes section at the end
    if notes:
        logger.debug(f"Adding {len(notes)} user notes to document...")
        doc.add_page_break()
        doc.add_heading("User Notes & Annotations", level=1)
        
        for note_idx, note in enumerate(notes, 1):
            note_para = doc.add_paragraph()
            
            # Add note number/location info
            location_str = ""
            if note.get('location'):
                loc = note['location']
                if 'page' in loc:
                    location_str = f" (Page {loc['page']})"
                elif 'location_start' in loc:
                    location_str = f" (Location {loc['location_start']})"
            
            # Add note header in bold
            run = note_para.add_run(f"Note {note_idx}{location_str}: ")
            run.bold = True
            
            # Add note text in italic
            run = note_para.add_run(note['text'])
            run.italic = True
            run.font.color.rgb = None  # Use default color
        
        logger.info(f"📝 Added {len(notes)} user notes at end of document")
    
    logger.debug("✓ Document creation complete")
    return doc, highlights_found

# --- MATCHER IMPLEMENTATIONS ---

def hybrid_matcher(book_text: str, highlights: List[str], threshold: float) -> Tuple[List[Tuple[int, int]], int]:
    """
    Hybrid matcher: Try regex first, then difflib for unmatched highlights.
    Uses progressively relaxed thresholds for difflib fallback.
    """
    logger.debug("Starting hybrid matching (regex → difflib fallback)...")
    found_spans = []
    matched_indices = set()
    
    # Phase 1: Regex matching (fast, exact) with FLEXIBLE spacing
    logger.info("Phase 1: Regex matching for exact matches...")
    highlights_sorted = sorted(enumerate(highlights), key=lambda x: len(x[1]), reverse=True)
    
    for idx, text in highlights_sorted:
        if idx % 10 == 0 and logger.level <= logging.DEBUG:
            logger.debug(f"Regex phase: {idx+1}/{len(highlights)}")
        
        # IMPROVED: Make spaces optional between words to handle spacing issues in conversion
        # Split on whitespace, escape each word, rejoin with flexible space pattern
        words = text.split()
        if len(words) > 1:
            # For multi-word highlights, allow 0-3 spaces between words
            escaped_words = [re.escape(word) for word in words]
            pattern = r'\s{0,3}'.join(escaped_words)
        else:
            # Single word - use standard escaping
            pattern = re.escape(text).replace(r'\ ', r'\s+')
        
        # Also handle various hyphen/dash characters
        pattern = pattern.replace(r'\-', r'[-\u2010-\u2015]')
        
        try:
            for match in re.finditer(pattern, book_text, re.IGNORECASE | re.UNICODE):
                start, end = match.span()
                if not any(max(start, s) < min(end, e) for s, e in found_spans):
                    found_spans.append((start, end))
                    matched_indices.add(idx)
                    break
        except (re.error, Exception) as e:
            if logger.level <= logging.DEBUG:
                logger.debug(f"Regex error for highlight {idx}: {e}")
            continue
    
    regex_count = len(matched_indices)
    logger.info(f"✓ Regex found {regex_count}/{len(highlights)} highlights ({regex_count/len(highlights)*100:.1f}%)")
    
    # Phase 2: Difflib fallback for unmatched highlights
    unmatched = [(idx, text) for idx, text in enumerate(highlights) if idx not in matched_indices]
    
    if not unmatched:
        logger.info("All highlights matched with regex!")
        return found_spans, regex_count
    
    logger.info(f"Phase 2: Difflib fallback for {len(unmatched)} unmatched highlights...")
    logger.debug("These may be nested/overlapping highlights or have spacing issues...")
    
    # Try progressively relaxed thresholds - LOWER thresholds for spacing issues
    thresholds = [0.95, 0.90, 0.85, 0.80, 0.75, 0.70, 0.65, threshold]
    thresholds = sorted(set(thresholds), reverse=True)
    
    # Track failures for verbose logging
    unmatched_details = {}
    
    for current_threshold in thresholds:
        if not unmatched:
            break
            
        logger.debug(f"Trying difflib with threshold {current_threshold:.2f} for {len(unmatched)} highlights...")
        newly_matched = []
        
        for idx, text in unmatched:
            if idx % 20 == 0 and logger.level <= logging.DEBUG:
                logger.debug(f"Difflib phase (t={current_threshold:.2f}): {len(newly_matched)} matched so far")
            
            normalized_highlight = re.sub(r'\s+', ' ', text).lower().strip()
            if not normalized_highlight or len(normalized_highlight) < 3:
                continue
            
            try:
                # Use difflib on ORIGINAL text (lowercased)
                matcher = SequenceMatcher(None, book_text.lower(), normalized_highlight)
                match = matcher.find_longest_match(0, len(book_text), 0, len(normalized_highlight))
                
                if match.size == 0:
                    unmatched_details[idx] = {
                        'text': text[:100] + '...' if len(text) > 100 else text,
                        'reason': 'No match found in text',
                        'best_score': 0.0
                    }
                    continue
                
                similarity = match.size / len(normalized_highlight)
                
                # Store best attempt for logging
                if idx not in unmatched_details or similarity > unmatched_details[idx].get('best_score', 0):
                    unmatched_details[idx] = {
                        'text': text[:100] + '...' if len(text) > 100 else text,
                        'reason': f'Best similarity: {similarity:.2%} (threshold: {current_threshold:.2%})',
                        'best_score': similarity,
                        'match_pos': match.a
                    }
                
                if similarity >= current_threshold:
                    # Get boundaries
                    start = match.a
                    end = match.a + match.size
                    
                    # Expand to full words
                    while start > 0 and book_text[start-1].isalnum():
                        start -= 1
                    while end < len(book_text) and book_text[end].isalnum():
                        end += 1
                    
                    # Try MORE FLEXIBLE regex in expanded region
                    search_start = max(0, start - 100)
                    search_end = min(len(book_text), end + 100)
                    search_region = book_text[search_start:search_end]
                    
                    # Build flexible pattern
                    words = text.split()
                    if len(words) > 1:
                        escaped_words = [re.escape(word) for word in words]
                        pattern = r'\s{0,3}'.join(escaped_words)
                    else:
                        pattern = re.escape(text).replace(r'\ ', r'\s+')
                    pattern = pattern.replace(r'\-', r'[-\u2010-\u2015]')
                    
                    found_exact = False
                    
                    try:
                        regex_matches = list(re.finditer(pattern, search_region, re.IGNORECASE | re.UNICODE))
                        
                        for regex_match in regex_matches:
                            exact_start = search_start + regex_match.start()
                            exact_end = search_start + regex_match.end()
                            
                            # For high-confidence matches (≥80%), ALLOW overlaps!
                            # Lower threshold because spacing issues reduce similarity
                            if similarity >= 0.80:
                                found_spans.append((exact_start, exact_end))
                                matched_indices.add(idx)
                                newly_matched.append(idx)
                                found_exact = True
                                if idx in unmatched_details:
                                    del unmatched_details[idx]
                                if logger.level <= logging.DEBUG:
                                    logger.debug(f"       ✓ Highlight {idx}: Added match at {exact_start}-{exact_end} (allowing overlap, sim={similarity:.2%})")
                                break
                            else:
                                # For lower confidence, still check for overlaps
                                has_overlap = any(max(exact_start, s) < min(exact_end, e) for s, e in found_spans)
                                if not has_overlap:
                                    found_spans.append((exact_start, exact_end))
                                    matched_indices.add(idx)
                                    newly_matched.append(idx)
                                    found_exact = True
                                    if idx in unmatched_details:
                                        del unmatched_details[idx]
                                    if logger.level <= logging.DEBUG:
                                        logger.debug(f"       ✓ Highlight {idx}: Added match at {exact_start}-{exact_end}")
                                    break
                    except re.error as e:
                        if logger.level <= logging.DEBUG:
                            logger.debug(f"Regex error: {e}")
                        pass
                    
                    # Fallback to fuzzy boundaries if exact regex failed
                    if not found_exact and similarity >= 0.80:  # Lower threshold
                        found_spans.append((start, end))
                        matched_indices.add(idx)
                        newly_matched.append(idx)
                        if idx in unmatched_details:
                            del unmatched_details[idx]
                        if logger.level <= logging.DEBUG:
                            logger.debug(f"       ✓ Highlight {idx}: Using fuzzy boundaries at {start}-{end} (allowing overlap, sim={similarity:.2%})")
                
            except Exception as e:
                logger.debug(f"Error matching highlight {idx}: {e}")
                unmatched_details[idx] = {
                    'text': text[:100] + '...' if len(text) > 100 else text,
                    'reason': f'Exception: {str(e)[:50]}',
                    'best_score': 0.0
                }
                continue
        
        # Remove newly matched from unmatched list
        unmatched = [item for item in unmatched if item[0] not in newly_matched]
        
        if newly_matched:
            logger.debug(f"✓ Threshold {current_threshold:.2f}: matched {len(newly_matched)} additional highlights")
    
    difflib_count = len(matched_indices) - regex_count
    total_count = len(matched_indices)
    
    logger.info(f"✓ Difflib fallback found {difflib_count} additional highlights")
    logger.info(f"📊 Total: {total_count}/{len(highlights)} highlights ({total_count/len(highlights)*100:.1f}%)")
    logger.debug(f"   Regex: {regex_count}, Difflib: {difflib_count}, Unmatched: {len(unmatched)}")
    
    if difflib_count > 0:
        logger.info(f"💡 Note: {difflib_count} highlights had spacing/formatting differences from conversion")
    
    # VERBOSE LOGGING for failed highlights
    if unmatched_details and logger.level <= logging.DEBUG:
        logger.debug("\n" + "="*70)
        logger.debug(f"❌ UNMATCHED HIGHLIGHTS DETAILS ({len(unmatched_details)} highlights):")
        logger.debug("="*70)
        for idx, details in sorted(unmatched_details.items()):
            logger.debug(f"\nHighlight {idx+1}:")
            logger.debug(f"  Text: {details['text']}")
            logger.debug(f"  Reason: {details['reason']}")
            if 'match_pos' in details:
                logger.debug(f"  Best match position: {details['match_pos']}")
                pos = details['match_pos']
                context_start = max(0, pos - 50)
                context_end = min(len(book_text), pos + 150)
                logger.debug(f"  Context: ...{book_text[context_start:context_end]}...")
    
    return found_spans, total_count

def vector_matcher(book_text: str, highlights: List[str], threshold: float) -> Tuple[List[Tuple[int, int]], int]:
    """Finds highlights using semantic vector similarity with improved confirmation."""
    logger.debug("Starting vector matching...")
    
    # Set environment variables BEFORE any imports
    import os
    os.environ['USE_TF'] = '0'
    os.environ['USE_TORCH'] = '1'
    os.environ['TRANSFORMERS_NO_TF'] = '1'
    os.environ['TF_CPP_MIN_LOG_LEVEL'] = '3'
    
    # Disable Triton if on Windows/Mac
    if sys.platform in ['win32', 'darwin']:
        logger.debug(f"Platform {sys.platform} detected, disabling Triton optimizations")
        os.environ['TRITON_DISABLE'] = '1'
    
    if not check_nlp_libraries():
        handle_error(
            "The 'vector' method requires PyTorch, sentence-transformers, and nltk.\n"
            "Install with: pip install torch sentence-transformers nltk\n"
            "\n"
            "If you have TensorFlow installed and getting errors, uninstall it:\n"
            "  pip uninstall tensorflow tf-keras\n"
            "\n"
            "Or use the 'hybrid' method instead (recommended):\n"
            "  python highlight.py --clippings ... -m hybrid"
        )
    
    try:
        import nltk
        from sentence_transformers import SentenceTransformer, util
        import torch
    except ImportError as e:
        handle_error(f"Failed to import required libraries: {e}")
    except Exception as e:
        handle_error(f"Error importing NLP libraries: {e}")
    
    logger.debug("Checking for NLTK punkt tokenizer...")
    try:
        nltk.data.find('tokenizers/punkt')
        logger.debug("✓ punkt tokenizer found")
    except LookupError:
        logger.info("Downloading NLTK 'punkt' tokenizer...")
        try:
            nltk.download('punkt', quiet=True)
            logger.debug("✓ punkt tokenizer downloaded")
        except Exception as e:
            logger.error(f"Failed to download punkt tokenizer: {e}")
            handle_error("Could not download required NLTK data")
    
    logger.info(f"Loading sentence transformer model: {NLP_MODEL}")
    try:
        model = SentenceTransformer(NLP_MODEL, device='cpu')
        logger.debug(f"✓ Model loaded on CPU")
    except Exception as e:
        logger.error(f"Failed to load model: {e}")
        logger.debug("Full traceback:", exc_info=True)
        handle_error(
            f"Could not load NLP model: {e}\n\n"
            "This might be due to TensorFlow conflicts. Try:\n"
            "  pip uninstall tensorflow tf-keras\n"
            "  pip install --upgrade sentence-transformers"
        )
    
    logger.info("Tokenizing book text and mapping sentence positions...")
    try:
        book_sentences_raw = nltk.sent_tokenize(book_text)
        logger.debug(f"✓ Book split into {len(book_sentences_raw)} sentences")

        # Pre-calculate all sentence positions
        sentence_positions = []
        book_sentences = []
        current_pos = 0
        for sent in book_sentences_raw:
            pos = book_text.find(sent, current_pos)
            if pos != -1:
                sentence_positions.append((pos, pos + len(sent)))
                book_sentences.append(sent)
                current_pos = pos + len(sent)
            else:
                logger.debug(f"Could not map sentence: {sent[:60]}...")

        logger.debug(f"✓ Successfully mapped {len(book_sentences)} sentences to character positions.")

    except Exception as e:
        handle_error(f"Failed during sentence tokenization: {e}")

    logger.info("Encoding book sentences...")
    book_embeddings = model.encode(
        book_sentences,
        convert_to_tensor=True,
        show_progress_bar=logger.level <= logging.INFO,
        normalize_embeddings=True,
        device='cpu'
    )

    logger.info("Encoding highlights...")
    highlight_embeddings = model.encode(
        highlights,
        convert_to_tensor=True,
        show_progress_bar=logger.level <= logging.INFO,
        normalize_embeddings=True,
        device='cpu'
    )

    logger.info("Computing similarity scores...")
    cosine_scores = util.cos_sim(highlight_embeddings, book_embeddings)
    
    found_spans = []
    unmatched_details = {}
    
    # Prepare normalized book text for difflib fallback
    normalized_book = re.sub(r'\s+', ' ', book_text).lower()
    
    logger.debug(f"Matching highlights with improved confirmation logic (threshold: {threshold})...")
    for i, text in enumerate(highlights):
        if i % 10 == 0 and logger.level <= logging.DEBUG:
            logger.debug(f"Processing highlight {i+1}/{len(highlights)}")
            
        try:
            # INCREASED CANDIDATES for better coverage
            top_k = min(20, len(book_sentences))  # Increased from 10 to 20
            top_scores, top_indices = torch.topk(cosine_scores[i], k=top_k)
            
            best_score = top_scores[0].item()
            if logger.level <= logging.DEBUG and (i < 5 or i % 50 == 0):
                logger.debug(f"Highlight {i+1}: best candidate score = {best_score:.3f}")

            match_found = False
            best_attempt = None
            
            # Loop through ALL top candidates
            for rank, (score_tensor, idx_tensor) in enumerate(zip(top_scores, top_indices)):
                if match_found:
                    break
                
                sent_idx = idx_tensor.item()
                sent_score = score_tensor.item()
                
                # Check confidence threshold
                if sent_score < threshold:
                    break
                
                sent_start, sent_end = sentence_positions[sent_idx]
                
                # MUCH LARGER BUFFER for better text capture
                buffer = max(len(text) * 2, 500)  # At least 500 chars or 2x highlight length
                context_start = max(0, sent_start - buffer)
                context_end = min(len(book_text), sent_end + buffer)
                search_span = book_text[context_start:context_end]
                
                if not best_attempt:
                    best_attempt = {
                        'score': sent_score,
                        'region': (context_start, context_end),
                        'sentence_text': book_sentences[sent_idx][:100]
                    }
                
                # Method 1: Try regex confirmation
                pattern = re.escape(text).replace(r'\ ', r'\s+').replace(r'\-', r'[-\u2010-\u2015]')
                try:
                    for match in re.finditer(pattern, search_span, re.IGNORECASE | re.UNICODE):
                        final_start = context_start + match.start()
                        final_end = context_start + match.end()
                        if not any(max(final_start, s) < min(final_end, e) for s, e in found_spans):
                            found_spans.append((final_start, final_end))
                            match_found = True
                            if i < 5 or (i % 50 == 0 and logger.level <= logging.DEBUG):
                                logger.debug(f"✓ Confirmed via regex in vector region at {final_start}-{final_end}")
                            break
                except re.error:
                    pass
                
                # Method 2: If regex failed, try difflib within the semantic region
                if not match_found and rank < 5:  # Only try difflib for top 5 candidates
                    normalized_highlight = re.sub(r'\s+', ' ', text).lower().strip()
                    normalized_region = re.sub(r'\s+', ' ', search_span).lower()
                    
                    matcher = SequenceMatcher(None, normalized_region, normalized_highlight)
                    match = matcher.find_longest_match(0, len(normalized_region), 0, len(normalized_highlight))
                    
                    if match.size > 0:
                        similarity = match.size / len(normalized_highlight)
                        
                        # Use lower threshold for vector-guided difflib (since we already have semantic match)
                        if similarity >= 0.75:  # More lenient than standalone difflib
                            # Map back to original text
                            final_start = context_start + match.a
                            final_end = min(final_start + len(text) + 50, len(book_text))
                            
                            # Adjust to word boundaries
                            while final_start > 0 and book_text[final_start-1].isalnum():
                                final_start -= 1
                            while final_end < len(book_text) and book_text[final_end].isalnum():
                                final_end += 1
                            
                            if not any(max(final_start, s) < min(final_end, e) for s, e in found_spans):
                                found_spans.append((final_start, final_end))
                                match_found = True
                                if i < 5 or (i % 50 == 0 and logger.level <= logging.DEBUG):
                                    logger.debug(f"✓ Confirmed via difflib (sim={similarity:.2%}) in vector region at {final_start}-{final_end}")
                                break

            if not match_found:
                unmatched_details[i] = {
                    'text': text[:100] + '...' if len(text) > 100 else text,
                    'best_score': best_score,
                    'best_sentence': best_attempt['sentence_text'] if best_attempt else 'N/A',
                    'reason': f"Semantic match found (score={best_score:.3f}) but text confirmation failed"
                }
                if logger.level <= logging.DEBUG and i < 10:
                    logger.debug(f"  ✗ Could not confirm highlight {i+1} despite semantic score {best_score:.3f}")

        except Exception as e:
            logger.debug(f"Error processing highlight {i}: {e}")
            unmatched_details[i] = {
                'text': text[:100] + '...' if len(text) > 100 else text,
                'best_score': 0.0,
                'reason': f'Exception: {str(e)[:50]}'
            }
            continue
    
    # VERBOSE LOGGING for failed highlights
    if unmatched_details and logger.level <= logging.DEBUG:
        logger.debug("\n" + "="*70)
        logger.debug(f"❌ VECTOR UNMATCHED HIGHLIGHTS ({len(unmatched_details)} highlights):")
        logger.debug("="*70)
        for idx, details in sorted(list(unmatched_details.items())[:20]):  # Show first 20
            logger.debug(f"\nHighlight {idx+1}:")
            logger.debug(f"  Text: {details['text']}")
            logger.debug(f"  Semantic score: {details['best_score']:.3f}")
            logger.debug(f"  Best sentence: {details.get('best_sentence', 'N/A')}")
            logger.debug(f"  Reason: {details['reason']}")
        if len(unmatched_details) > 20:
            logger.debug(f"\n... and {len(unmatched_details) - 20} more unmatched highlights")
            
    logger.debug(f"✓ Vector matching complete. Found {len(found_spans)} matches")
    return found_spans, len(found_spans)

def regex_matcher(book_text: str, highlights: List[str]) -> Tuple[List[Tuple[int, int]], int]:
    """Finds highlights using flexible regular expressions."""
    logger.debug("Starting regex matching...")
    found_spans = []
    
    highlights_sorted = sorted(highlights, key=len, reverse=True)
    logger.debug(f"Processing {len(highlights_sorted)} highlights (longest first)")
    
    for idx, text in enumerate(highlights_sorted):
        if idx % 10 == 0 and logger.level <= logging.DEBUG:
            logger.debug(f"Processing highlight {idx+1}/{len(highlights)}")
        
        if idx == 0:
            logger.debug(f"First highlight (length {len(text)}): {text[:50]}...")
        
        pattern = re.escape(text).replace(r'\ ', r'\s+').replace(r'\-', r'[-\u2010-\u2015]')
        try:
            for match in re.finditer(pattern, book_text, re.IGNORECASE | re.UNICODE):
                start, end = match.span()
                if not any(max(start, s) < min(end, e) for s, e in found_spans):
                    found_spans.append((start, end))
                    if idx < 5:
                        logger.debug(f"✓ Match found at position {start}-{end}")
                    break
        except re.error as e:
            logger.debug(f"Regex error for pattern: {e}")
            continue
        except Exception as e:
            logger.debug(f"Unexpected error in regex matching: {e}")
            continue
    
    logger.debug(f"✓ Regex matching complete. Found {len(found_spans)} matches")
    return found_spans, len(found_spans)

def difflib_matcher(book_text: str, highlights: List[str], threshold: float) -> Tuple[List[Tuple[int, int]], int]:
    """Improved difflib matcher with better performance and accuracy."""
    logger.debug(f"Starting difflib matching with threshold {threshold}...")
    found_spans = []
    
    logger.debug("Normalizing book text...")
    normalized_book = re.sub(r'\s+', ' ', book_text).lower()
    logger.debug(f"Normalized text length: {len(normalized_book)}")
    
    highlights_sorted = sorted(enumerate(highlights), key=lambda x: len(x[1]), reverse=True)
    logger.debug(f"Processing {len(highlights_sorted)} highlights")
    
    matched_count = 0
    
    for list_idx, (orig_idx, text) in enumerate(highlights_sorted):
        if list_idx % 10 == 0 and logger.level <= logging.DEBUG:
            logger.debug(f"Processing highlight {list_idx+1}/{len(highlights)}")
        
        normalized_highlight = re.sub(r'\s+', ' ', text).lower().strip()
        if not normalized_highlight or len(normalized_highlight) < MIN_HIGHLIGHT_LENGTH:
            continue

        try:
            # Use SequenceMatcher for fuzzy matching
            matcher = SequenceMatcher(None, normalized_book, normalized_highlight, autojunk=False)
            match = matcher.find_longest_match(0, len(normalized_book), 0, len(normalized_highlight))
            
            if match.size == 0:
                continue
            
            similarity = match.size / len(normalized_highlight)
            
            if list_idx < 5:
                logger.debug(f"Highlight {list_idx}: similarity = {similarity:.3f}, match_size = {match.size}/{len(normalized_highlight)}")
            
            if similarity >= threshold:
                # Map normalized position back to original text position
                # Find corresponding position in original text
                start = match.a
                
                # Calculate end position with buffer for formatting differences
                end = min(start + int(len(text) * 1.2), len(book_text))
                
                # Refine boundaries to avoid cutting words
                while start > 0 and book_text[start-1:start].strip() == '' and book_text[start-1].isalnum():
                    start -= 1
                while end < len(book_text) and book_text[end:end+1].strip() == '' and book_text[end].isalnum():
                    end += 1
                
                # Check for overlap
                if not any(max(start, s) < min(end, e) for s, e in found_spans):
                    found_spans.append((start, end))
                    matched_count += 1
                    if list_idx < 5:
                        logger.debug(f"✓ Match found with {similarity:.2%} similarity at position {start}-{end}")
        except Exception as e:
            logger.debug(f"Error matching highlight {orig_idx}: {e}")
            continue
    
    logger.debug(f"✓ Difflib matching complete. Found {matched_count} matches")
    return found_spans, matched_count

# --- SINGLE BOOK PROCESSING ---

def process_single_book(
    ebook_path: str,
    all_highlights: List[Dict[str, str]],
    all_notes: List[Dict[str, str]],
    converter_path: str,
    output_format: str,
    method: str = 'diff',
    output_path: Optional[str] = None,
    keep_html: bool = False,
    similarity_threshold: float = DEFAULT_SIMILARITY_THRESHOLD,
    vector_threshold: float = DEFAULT_VECTOR_CONFIDENCE_THRESHOLD,
    compare_mode: bool = False
) -> Dict:
    """
    Process a single book and return results.
    Returns dict with: success, output_path, highlights_found, highlights_total, error
    """
    logger.debug(f"Starting process_single_book for: {ebook_path}")
    result = {
        'success': False,
        'ebook_path': str(ebook_path),
        'output_path': None,
        'highlights_found': 0,
        'highlights_total': 0,
        'notes_count': 0,
        'error': None
    }
    
    html_file = None
    intermediate_file = None
    extract_dir = None
    
    try:
        # Convert ebook
        logger.debug("Converting ebook to HTML...")
        html_file, intermediate_file, extract_dir = convert_ebook_to_html(
            str(ebook_path), converter_path, output_format
        )
        
        # Extract title
        logger.debug(f"Reading HTML file for title extraction: {html_file}")
        try:
            with open(html_file, 'r', encoding='utf-8') as f:
                soup = BeautifulSoup(f, 'html.parser')
                doc_title = soup.title.string.strip() if soup.title and soup.title.string else os.path.basename(ebook_path)
        except Exception as e:
            logger.warning(f"Could not extract title from HTML: {e}")
            doc_title = os.path.basename(ebook_path)
        
        logger.debug(f"Document title: '{doc_title}'")
        
        # Find highlights and notes
        logger.debug("Finding highlights and notes for this book...")
        relevant_highlights, relevant_notes = find_book_highlights_and_notes(doc_title, all_highlights, all_notes)
        
        if not relevant_highlights and not relevant_notes:
            result['error'] = "No highlights or notes found"
            logger.warning("No highlights or notes found for this book")
            return result
        
        result['highlights_total'] = len(relevant_highlights)
        result['notes_count'] = len(relevant_notes)
        logger.debug(f"Found {len(relevant_highlights)} highlights and {len(relevant_notes)} notes to process")
        
        # --- COMPARE MODE ---
        if compare_mode:
            logger.debug("Comparison mode activated")
            logger.info("\n" + "=" * 70)
            logger.info("⚖️  Running Comparison Mode")
            logger.info("=" * 70)
            
            results = {}
            
            # Test Regex
            try:
                logger.info(f"\n🧪 Testing Regex method...")
                _, count = create_highlighted_docx(
                    html_file, relevant_highlights, relevant_notes, doc_title,
                    lambda txt, h, t: regex_matcher(txt, h),
                    threshold=0.0
                )
                results["Regex"] = f"{count}/{len(relevant_highlights)} ({count/len(relevant_highlights)*100:.1f}%)"
                logger.debug(f"Regex result: {results['Regex']}")
            except Exception as e:
                logger.error(f"Regex method failed: {e}")
                logger.debug(f"Regex exception details:", exc_info=True)
                results["Regex"] = f"Failed ({type(e).__name__})"
            
            # Test Difflib
            try:
                logger.info(f"\n🧪 Testing Difflib method...")
                _, count = create_highlighted_docx(
                    html_file, relevant_highlights, relevant_notes, doc_title,
                    difflib_matcher,
                    similarity_threshold
                )
                results["Difflib"] = f"{count}/{len(relevant_highlights)} ({count/len(relevant_highlights)*100:.1f}%)"
                logger.debug(f"Difflib result: {results['Difflib']}")
            except Exception as e:
                logger.error(f"Difflib method failed: {e}")
                logger.debug(f"Difflib exception details:", exc_info=True)
                results["Difflib"] = f"Failed ({type(e).__name__})"
            
            # Test Vector
            try:
                logger.info(f"\n🧪 Testing Vector method...")
                _, count = create_highlighted_docx(
                    html_file, relevant_highlights, relevant_notes, doc_title,
                    vector_matcher,
                    vector_threshold
                )
                results["Vector"] = f"{count}/{len(relevant_highlights)} ({count/len(relevant_highlights)*100:.1f}%)"
                logger.debug(f"Vector result: {results['Vector']}")
            except Exception as e:
                logger.error(f"Vector method failed: {e}")
                logger.debug(f"Vector exception details:", exc_info=True)
                results["Vector"] = f"Failed ({type(e).__name__})"

            print("\n" + "=" * 70)
            print("📊 Comparison Results")
            print("=" * 70)
            for name, res in results.items():
                print(f"  {name:<10}: {res}")
            print("=" * 70)
            
            result['success'] = True
            result['compare_results'] = results
            return result
        
        # --- SINGLE METHOD MODE ---
        logger.debug(f"Single method mode: {method}")
        
        # Select matcher
        if method == 'regex':
            selected_matcher = lambda txt, h, t: regex_matcher(txt, h)
            threshold = 0.0
            logger.debug("Using regex matcher (no threshold)")
        elif method == 'diff':
            selected_matcher = difflib_matcher
            threshold = similarity_threshold
            logger.debug(f"Using difflib matcher (threshold: {threshold})")
        elif method == 'hybrid':
            selected_matcher = hybrid_matcher
            threshold = similarity_threshold
            logger.debug(f"Using hybrid matcher (regex + difflib fallback, threshold: {threshold})")
        else:  # vector
            selected_matcher = vector_matcher
            threshold = vector_threshold
            logger.debug(f"Using vector matcher (threshold: {threshold})")

        logger.debug("Creating highlighted document...")
        doc, found_count = create_highlighted_docx(
            html_file,
            relevant_highlights,
            relevant_notes,
            doc_title,
            selected_matcher,
            threshold
        )
        
        result['highlights_found'] = found_count
        logger.debug(f"Found {found_count} highlights")
        
        # Save document
        if not output_path:
            output_path = os.path.splitext(ebook_path)[0] + '_highlighted.docx'
        
        logger.debug(f"Output path: {output_path}")
        logger.info(f"💾 Saving document to: {output_path}")
        
        try:
            doc.save(output_path)
            logger.debug("✓ Document saved successfully")
        except Exception as e:
            logger.error(f"Failed to save document: {e}")
            logger.debug("Full traceback:", exc_info=True)
            result['error'] = f"Failed to save document: {e}"
            return result
        
        result['success'] = True
        result['output_path'] = output_path
        
    except Exception as e:
        logger.error(f"Error processing book: {e}")
        logger.debug("Full exception:", exc_info=True)
        result['error'] = str(e)
    
    finally:
        # Cleanup
        if not keep_html:
            logger.debug("Cleaning up intermediate files...")
            if html_file and os.path.exists(html_file):
                try:
                    os.remove(html_file)
                    logger.debug(f"✓ Cleaned up: {os.path.basename(html_file)}")
                except Exception as e:
                    logger.debug(f"Failed to clean up {html_file}: {e}")
            
            if extract_dir and os.path.exists(extract_dir):
                try:
                    shutil.rmtree(extract_dir)
                    logger.debug(f"✓ Cleaned up extracted directory: {extract_dir}")
                except Exception as e:
                    logger.debug(f"Failed to clean up {extract_dir}: {e}")
            
            if intermediate_file and os.path.exists(intermediate_file):
                try:
                    os.remove(intermediate_file)
                    logger.debug(f"✓ Cleaned up: {os.path.basename(intermediate_file)}")
                except Exception as e:
                    logger.debug(f"Failed to clean up {intermediate_file}: {e}")
        else:
            logger.info(f"Keeping intermediate files (--preserve-html):")
            if html_file:
                logger.info(f"  - HTML: {html_file}")
            if intermediate_file:
                logger.info(f"  - Archive: {intermediate_file}")
            if extract_dir:
                logger.info(f"  - Extract dir: {extract_dir}")
    
    return result

# --- MAIN EXECUTION ---

def main():
    parser = argparse.ArgumentParser(
        description="Generate Word documents with Kindle highlights - LIBRARY BATCH MODE BY DEFAULT",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  Default: Process ALL books from library with 95%+ match confidence:
    python kindle_highlighter.py --clippings "My Clippings.txt" -v
  
  Process only the best matching book from library:
    python kindle_highlighter.py --clippings "My Clippings.txt" --no-batch -v
  
  Process specific ebook file:
    python kindle_highlighter.py --ebook book.mobi --clippings "My Clippings.txt" -v
  
  Batch mode with HTML files preserved for debugging:
    python kindle_highlighter.py --clippings "My Clippings.txt" --preserve-html -vv
  
  Compare all methods on a specific file:
    python kindle_highlighter.py --ebook book.mobi --clippings "My Clippings.txt" --compare -v
  
  List all books in library:
    python kindle_highlighter.py --list-books

Note: Library batch mode is enabled by default. Use --no-batch to process only 
      the single best match, or --ebook to process a specific file.
        """
    )
    
    # Mode selection
    mode_group = parser.add_mutually_exclusive_group()
    mode_group.add_argument("--list-books", action="store_true", 
                           help="List all books in Calibre Library and exit.")
    
    # Direct file mode
    direct_group = parser.add_argument_group('Direct File Mode')
    direct_group.add_argument("--ebook", dest="ebook_file", 
                             help="Path to a single ebook file (overrides library mode).")
    direct_group.add_argument("--clippings", dest="clippings_file", 
                             help="Path to 'My Clippings.txt' (required).")
    
    # Library options (defaults enabled)
    library_group = parser.add_argument_group('Library Options')
    library_group.add_argument("--library-path", 
                              help=f"Path to Calibre Library (default: {DEFAULT_CALIBRE_LIBRARY})")
    library_group.add_argument("--no-batch", action="store_true",
                              help="Process only the single best match instead of all books with 95%+ confidence")
    
    # Converter options
    converter_group = parser.add_argument_group('Converter Options')
    converter_group.add_argument("--calibre-path", 
                                help="Path to Calibre's ebook-convert binary")
    converter_group.add_argument("--try-native", action="store_true", 
                                help="Try native Python conversion if converter fails.")
    
    # Output options
    output_group = parser.add_argument_group('Output Options')
    output_group.add_argument("-o", "--output", 
                             help="Output path (single file mode only)")
    output_group.add_argument("--keep-html", action="store_true", 
                             help="Keep intermediate HTML files")
    output_group.add_argument("--preserve-html", action="store_true",
                             help="Same as --keep-html (preserve HTML files for inspection)")
    
    # Matching options
    matching_group = parser.add_argument_group('Matching Options')
    matching_group.add_argument(
        "-m", "--method",
        choices=['regex', 'diff', 'hybrid', 'vector'],
        default='hybrid',  # Changed from 'diff' to 'hybrid'
        help="Matching method: regex (fast), diff (balanced), hybrid (best, default), vector (semantic)"
    )
    matching_group.add_argument("--compare", action="store_true", 
                               help="Run all methods and compare results (single file only)")
    matching_group.add_argument(
        "--similarity-threshold",
        type=float,
        default=DEFAULT_SIMILARITY_THRESHOLD,
        help=f"Similarity threshold for difflib (default: {DEFAULT_SIMILARITY_THRESHOLD})"
    )
    matching_group.add_argument(
        "--vector-threshold",
        type=float,
        default=DEFAULT_VECTOR_CONFIDENCE_THRESHOLD,
        help=f"Confidence threshold for vector (default: {DEFAULT_VECTOR_CONFIDENCE_THRESHOLD})"
    )
    
    # General options
    parser.add_argument(
        "-v", "--verbose",
        action="count",
        default=0,
        help="Increase verbosity (-v for info, -vv for debug)"
    )
    
    args = parser.parse_args()
    
    # Setup logging
    setup_logging(args.verbose)
    
    logger.info("=" * 70)
    logger.info("📚 Kindle Ebook Highlighter - COMPLETE VERSION")
    logger.info("=" * 70)
    logger.debug(f"Command line arguments: {vars(args)}")
    logger.debug(f"Python version: {sys.version}")
    logger.debug(f"Platform: {sys.platform}")
    
    # Handle preserve-html alias
    if args.preserve_html:
        args.keep_html = True
        logger.debug("--preserve-html flag set, enabling --keep-html")
    
    # --- LIST BOOKS MODE ---
    if args.list_books:
        logger.debug("List books mode activated")
        library_path = find_calibre_library(args.library_path)
        books = discover_books(library_path)
        
        print("\n" + "=" * 70)
        print(f"Books in Calibre Library ({library_path})")
        print("=" * 70)
        
        for book in books:
            print(f"📖 {book['title']}")
            print(f"   Author: {book['author']}")
            print(f"   Path: {book['path']}")
            print()
        
        print(f"Total: {len(books)} books")
        sys.exit(0)
    
    # Check dependencies
    logger.info("\n🔍 Checking dependencies...")
    deps = check_dependencies(args.calibre_path)
    
    required = ['docx', 'bs4', 'thefuzz', 'converter']
    missing_required = [k for k in required if not deps[k]]
    
    if missing_required:
        handle_error(f"Missing required dependencies: {', '.join(missing_required)}")
    
    if args.try_native and not (deps['mobi'] or deps['ebooklib']):
        logger.warning("--try-native flag used, but 'mobi' and 'ebooklib' libraries are not installed.")
        logger.warning("Install them: pip install mobi ebooklib")
    
    logger.info("✓ All required dependencies found\n")
    
    # --- CORE LOGIC ---
    try:
        logger.debug("Starting main processing...")
        
        # Get clippings path
        clippings_path = args.clippings_file
        
        # If no ebook file specified, use library mode by default
        if not args.ebook_file and not clippings_path:
            logger.debug("No ebook or clippings specified, checking for defaults...")
            clippings_default = os.path.expanduser('~/Documents/My Clippings.txt')
            if os.path.exists(clippings_default):
                clippings_path = clippings_default
                logger.info(f"Using default clippings: {clippings_path}")
        
        if not clippings_path:
            handle_error("A clippings file must be provided with --clippings.")
        
        logger.debug(f"Clippings path: {clippings_path}")
        
        # Parse clippings
        all_highlights, all_notes = parse_clippings(clippings_path)
        
        # Get converter
        converter_path, output_format = check_converter_available(args.calibre_path)
        
        if converter_path:
            logger.debug(f"Converter available: {converter_path}")
        elif args.try_native:
            logger.warning("No standard converter found. Will try native conversion.")
        else:
            handle_error("No converter available. Install Calibre or use --try-native")
        
        # --- LIBRARY MODE (DEFAULT) ---
        if not args.ebook_file:
            logger.debug("Library mode: processing books from Calibre library")
            library_path = find_calibre_library(args.library_path)
            books = discover_books(library_path)
            
            # Match all books - UPDATED
            matches = match_books_to_clippings(books, all_highlights, all_notes, MIN_TITLE_MATCH_SCORE)
            
            if not matches:
                handle_error("No matching books found in library!")
            
            # Show matches
            logger.info("\n" + "=" * 80)
            logger.info("Book Matching Results:")
            logger.info("=" * 80)
            for idx, m in enumerate(matches[:20], 1):
                logger.info(f"{idx:2d}. Score {m['score']:3d}% | {m['highlight_count']:4d} highlights | '{m['book']['title']}'")
                logger.info(f"     Clipping title: '{m['clipping_title']}'")
            if len(matches) > 20:
                logger.info(f"     ... and {len(matches) - 20} more books")
            logger.info("=" * 80)
            
            # Determine which books to process (BATCH BY DEFAULT)
            if args.no_batch:
                # Process only the best match
                if matches[0]['score'] >= AUTO_SELECT_THRESHOLD:
                    books_to_process = [matches[0]]
                    logger.info(f"\n✓ Processing single best match: '{matches[0]['book']['title']}' ({matches[0]['score']}%)")
                    high_conf_count = len([m for m in matches if m['score'] >= AUTO_SELECT_THRESHOLD])
                    if high_conf_count > 1:
                        logger.info(f"💡 Tip: Remove --no-batch to process all {high_conf_count} books with {AUTO_SELECT_THRESHOLD}%+ match")
                else:
                    logger.error(f"\n❌ Best match only {matches[0]['score']}% (need {AUTO_SELECT_THRESHOLD}%+)")
                    logger.error(f"   Book: '{matches[0]['book']['title']}'")
                    logger.error("   Try removing --no-batch, or use --ebook for specific file")
                    sys.exit(1)
            else:
                # DEFAULT: Process ALL books with 95%+ match (BATCH MODE)
                books_to_process = [m for m in matches if m['score'] >= AUTO_SELECT_THRESHOLD]
                logger.info(f"\n📦 BATCH MODE (default): Processing {len(books_to_process)} books with {AUTO_SELECT_THRESHOLD}%+ match confidence")
                if len(books_to_process) == 0:
                    logger.error(f"No books found with {AUTO_SELECT_THRESHOLD}%+ match confidence")
                    logger.error(f"Best match was {matches[0]['score']}%: '{matches[0]['book']['title']}'")
                    sys.exit(1)
            
            results = []
            for idx, match in enumerate(books_to_process, 1):
                book = match['book']
                logger.info(f"\n{'='*70}")
                logger.info(f"Processing {idx}/{len(books_to_process)}: {book['title']}")
                logger.info(f"Match: {match['clipping_title']} ({match['score']}%, {match['highlight_count']} items)")
                logger.info(f"{'='*70}")
                
                result = process_single_book(
                    book['path'],
                    all_highlights,  # UPDATED
                    all_notes,       # UPDATED
                    converter_path,
                    output_format,
                    method=args.method,
                    keep_html=args.keep_html,
                    similarity_threshold=args.similarity_threshold,
                    vector_threshold=args.vector_threshold,
                    compare_mode=False
                )
                
                results.append({
                    'book_title': book['title'],
                    'match_score': match['score'],
                    **result
                })
            
            # Summary - UPDATED to show notes
            print("\n" + "=" * 70)
            print("📊 BATCH PROCESSING SUMMARY")
            print("=" * 70)
            
            successful = [r for r in results if r['success']]
            failed = [r for r in results if not r['success']]
            
            print(f"Total books processed: {len(results)}")
            print(f"✅ Successful: {len(successful)}")
            print(f"❌ Failed: {len(failed)}")
            print()
            
            if successful:
                print("✅ Successful:")
                for r in successful:
                    pct = (r['highlights_found']/r['highlights_total']*100) if r['highlights_total'] > 0 else 0
                    print(f"   • {r['book_title']}")
                    print(f"     Found {r['highlights_found']}/{r['highlights_total']} highlights ({pct:.1f}%)")
                    if r.get('notes_count', 0) > 0:
                        print(f"     Added {r['notes_count']} user notes")
                    print(f"     Saved to: {r['output_path']}")
                print()
            
            if failed:
                print("❌ Failed:")
                for r in failed:
                    print(f"   • {r['book_title']}: {r['error']}")
                print()
            
            print("=" * 70)
        
        # --- SINGLE FILE MODE ---
        else:
            ebook_path = args.ebook_file
            logger.debug(f"Single file mode: processing {ebook_path}")
            
            result = process_single_book(
                ebook_path,
                all_highlights, 
                all_notes, 
                converter_path,
                output_format,
                method=args.method,
                output_path=args.output,
                keep_html=args.keep_html,
                similarity_threshold=args.similarity_threshold,
                vector_threshold=args.vector_threshold,
                compare_mode=args.compare
            )
            
            if args.compare:
                logger.debug("Compare mode results displayed")
            elif result['success']:
                pct = (result['highlights_found']/result['highlights_total']*100) if result['highlights_total'] > 0 else 0
                print("\n" + "=" * 70)
                print(f"✅ Success! Found {result['highlights_found']}/{result['highlights_total']} highlights ({pct:.1f}%)")
                if result.get('notes_count', 0) > 0:
                    print(f"📝 Added {result['notes_count']} user notes")
                print(f"📄 Document saved to: {result['output_path']}")
                print("=" * 70)
            else:
                print("\n" + "=" * 70)
                print(f"❌ Failed: {result['error']}")
                print("=" * 70)
                sys.exit(1)

    except KeyboardInterrupt:
        logger.warning("\n\n⚠️  Operation cancelled by user")
        sys.exit(1)
    except Exception as e:
        logger.exception("\n❌ Unexpected error occurred:")
        sys.exit(1)

if __name__ == "__main__":
    main()